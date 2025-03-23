import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def convert_md_to_docx(md_file, output_file):
    """Convert markdown file to Word document with proper formatting."""
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert markdown to HTML
    html_content = markdown.markdown(md_content)
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Process HTML content
    current_heading_level = 0
    for line in html_content.split('\n'):
        if line.startswith('<h'):
            # Handle headings
            level = int(line[2])
            text = line[4:-5]  # Remove HTML tags
            heading = doc.add_heading(text, level=level)
            heading.style.font.name = 'Calibri'
            heading.style.font.size = Pt(14 - level)  # Larger font for headings
            current_heading_level = level
        elif line.startswith('<p>'):
            # Handle paragraphs
            text = line[3:-4]  # Remove HTML tags
            p = doc.add_paragraph(text)
            p.style.font.name = 'Calibri'
            p.style.font.size = Pt(11)
        elif line.startswith('<pre>'):
            # Handle code blocks
            code = line[5:-6]  # Remove HTML tags
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif line.startswith('<ul>'):
            # Handle lists
            items = line[4:-5].split('</li>')
            for item in items:
                if item.strip():
                    text = item.replace('<li>', '').strip()
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(text)
    
    # Save the document
    doc.save(output_file)
    print(f"Document saved as {output_file}")

if __name__ == "__main__":
    # Convert markdown to Word document
    md_file = 'project_documentation.md'
    output_file = 'Device_Health_Prediction_Model_Documentation.docx'
    convert_md_to_docx(md_file, output_file) 